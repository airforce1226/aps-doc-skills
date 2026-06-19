# -*- coding: utf-8 -*-
"""Build a styled Korean .docx from a JSON spec.

Usage:
    python build_doc.py spec.json output.docx

Spec format (UTF-8 JSON):
{
  "title": "보고서 제목",
  "subtitle": "부제 (선택)",
  "meta": "작성일 ... | 작성 ...  (선택)",
  "accent": "1F6F43",          # 강조색 hex (선택, 기본 녹색)
  "font": "맑은 고딕",          # 본문 글꼴 (선택)
  "sections": [
    {
      "h1": "1. 사업 개요",
      "blocks": [
        {"bullet": "내용", "label": "사업명"},      # label 선택
        {"para": "문단 내용", "grey": true},        # grey/italic/size 선택
        {"h2": "1.1 소제목"},
        {"table": {
            "headers": ["열1", "열2"],
            "rows": [["a", "b"], ["c", "TBD"]],
            "widths": [4.0, 8.0]                     # cm, 선택
        }}
      ]
    }
  ],
  "footnote": "※ ... (선택)"
}

값이 "TBD" 인 셀/문구는 회색 이탤릭(미정 표시)으로 렌더링된다.
"""
import json
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GREY = RGBColor(0x59, 0x59, 0x59)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TBD = "TBD"


def build(spec, out_path):
    accent = RGBColor.from_string(spec.get("accent", "1F6F43"))
    accent_fill = spec.get("accent", "1F6F43")
    font_name = spec.get("font", "맑은 고딕")

    doc = Document()
    base = doc.styles["Normal"]
    base.font.name = font_name
    base.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    base.font.size = Pt(10.5)

    def style_run(run, bold=False, size=10.5, color=None, italic=False):
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        run.font.bold = bold
        run.font.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color

    def set_cell_bg(cell, color):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), color)
        tcPr.append(shd)

    def add_table(spec_t):
        headers = spec_t["headers"]
        rows = spec_t.get("rows", [])
        widths = spec_t.get("widths")
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, htext in enumerate(headers):
            cell = t.rows[0].cells[i]
            set_cell_bg(cell, accent_fill)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style_run(p.add_run(str(htext)), bold=True, color=WHITE)
        for row in rows:
            cells = t.add_row().cells
            for i, val in enumerate(row):
                p = cells[i].paragraphs[0]
                text = "" if val is None else str(val)
                is_tbd = text.strip() == TBD or text.startswith("(미정")
                r = p.add_run("(미정 — 추후 확정)" if text.strip() == TBD else text)
                style_run(r, size=9.5, color=GREY if is_tbd else None, italic=is_tbd)
        if widths:
            for row in t.rows:
                for i, w in enumerate(widths):
                    row.cells[i].width = Cm(w)
        return t

    # ---- cover ----
    if spec.get("title"):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_run(p.add_run(spec["title"]), bold=True, size=20, color=accent)
    if spec.get("subtitle"):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_run(p.add_run(spec["subtitle"]), size=11, color=GREY)
    if spec.get("meta"):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_run(p.add_run(spec["meta"]), size=10, color=GREY)
    doc.add_paragraph()

    # ---- sections ----
    for sec in spec.get("sections", []):
        if sec.get("h1"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
            style_run(p.add_run(sec["h1"]), bold=True, size=14, color=accent)
        for block in sec.get("blocks", []):
            if "h2" in block:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(2)
                style_run(p.add_run(block["h2"]), bold=True, size=11.5)
            elif "bullet" in block:
                p = doc.add_paragraph(style="List Bullet")
                if block.get("label"):
                    style_run(p.add_run(block["label"] + " : "), bold=True)
                style_run(p.add_run(block["bullet"]))
            elif "para" in block:
                p = doc.add_paragraph()
                style_run(
                    p.add_run(block["para"]),
                    size=block.get("size", 10.5),
                    color=GREY if block.get("grey") else None,
                    italic=block.get("italic", False),
                )
            elif "table" in block:
                add_table(block["table"])

    if spec.get("footnote"):
        doc.add_paragraph()
        p = doc.add_paragraph()
        style_run(p.add_run(spec["footnote"]), size=9, color=GREY, italic=True)

    doc.save(out_path)
    return out_path


def main():
    if len(sys.argv) < 3:
        print("Usage: python build_doc.py spec.json output.docx", file=sys.stderr)
        sys.exit(2)
    with open(sys.argv[1], "r", encoding="utf-8") as f:
        spec = json.load(f)
    out = build(spec, sys.argv[2])
    print("Saved:", out)


if __name__ == "__main__":
    main()
